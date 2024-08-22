import os
from docx import Document
import logging
import nltk
from nltk.tokenize import sent_tokenize
import sys

# Download the NLTK data files
nltk.download('punkt')

# Configure logging with UTF-8 encoding
log_file = 'app.log'
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler(log_file, encoding='utf-8'),  # Set encoding to UTF-8
        logging.StreamHandler(sys.stdout)  # Use sys.stdout directly for console output
    ]
)

def split_by_sentences(text):
    """Split text into sentences using NLTK."""
    return sent_tokenize(text)

def read_text_file(file_path):
    """Read text file and split content into sentences."""
    content = []
    with open(file_path, 'r', encoding='utf-8') as file:
        text = file.read()
        sentences = split_by_sentences(text)
        content.extend(sentences)
    return content

def collect_and_save_to_docx(base_dir, output_file):
    """Collect all text files from a directory and save the content to a Word document."""
    doc = Document()

    def add_text_from_file(directory):
        for file in os.listdir(directory):
            file_path = os.path.join(directory, file)
            if os.path.isfile(file_path) and file.endswith('.txt'):
                with open(file_path, 'r', encoding='utf-8') as f:
                    doc.add_heading(file, level=1)
                    file_content = f.read()
                    sentences = split_by_sentences(file_content)
                    for sentence in sentences:
                        doc.add_paragraph(sentence)
                    doc.add_page_break()
            else:
                logging.info(f"Skipping non-text file: {file}")

    def dir_or_not(base_dir):
        for item in os.listdir(base_dir):
            item_path = os.path.join(base_dir, item)
            if os.path.isdir(item_path):
                return True
        return False

    has_subdir = dir_or_not(base_dir)

    if has_subdir:
        for subdir in os.listdir(base_dir):
            subdir_path = os.path.join(base_dir, subdir)
            if os.path.isdir(subdir_path):
                add_text_from_file(subdir_path)
    else:
        add_text_from_file(base_dir)

    output_folder = os.path.join(base_dir, output_file)
    doc.save(output_folder)
    logging.info(f"All text from txt files in {base_dir} has been saved to {output_file}.")

def inttt(base_dir, output_file):
    collect_and_save_to_docx(base_dir, output_file)

if __name__ == "__main__":
    base_dir = sys.argv[1]
    output_file = sys.argv[2]
    try:
        inttt(base_dir, output_file)
    except Exception as e:
        logging.error(f"An error occurred: {str(e)}")
